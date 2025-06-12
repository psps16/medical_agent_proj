import markdown
import xml.etree.ElementTree as ET # Import ElementTree for parsing HTML
from docx import Document
from docx.shared import Pt # For font size

class ReportToolDocx:
    """
    A class to convert Markdown content to a formatted DOCX file.
    """

    @staticmethod
    def process_markdown_element(element, document):
        """
        Recursively process markdown element tree (parsed from HTML) and add elements to a python-docx Document.
        Handles basic elements like headings, paragraphs, bold, italics, lists, and line breaks.
        This is a static method as it operates solely on its inputs and doesn't require instance state.
        """
        # Handle block-level elements
        if element.tag == 'h1':
            heading = document.add_heading(level=1)
            # Process inline elements within the heading
            ReportToolDocx._process_inline_elements(element, heading)
        elif element.tag == 'h2':
            heading = document.add_heading(level=2)
            ReportToolDocx._process_inline_elements(element, heading)
        elif element.tag == 'h3':
            heading = document.add_heading(level=3)
            ReportToolDocx._process_inline_elements(element, heading)
        elif element.tag == 'p':
            paragraph = document.add_paragraph()
            ReportToolDocx._process_inline_elements(element, paragraph)
        elif element.tag == 'ul' or element.tag == 'ol':
            # Process list items
            for li in element.findall('li'):
                # python-docx handles list formatting automatically when adding list items
                # We need to add the list item text to a paragraph and then apply the list style
                list_paragraph = document.add_paragraph(style='List Bullet' if element.tag == 'ul' else 'List Number')
                ReportToolDocx._process_inline_elements(li, list_paragraph)

        elif element.tag == 'hr':
            # Horizontal rule - python-docx doesn't have a direct HR element.
            # We can simulate it with a paragraph containing underscores or just add space.
            # Adding a simple line for now.
            document.add_paragraph("___") # Simple visual separator
            document.add_paragraph() # Add an empty line for spacing

        # Add handling for other tags like code blocks, blockquotes, tables, etc. if needed
        # For now, if a tag is not explicitly handled as a block, process its children
        else:
            # Process children of unhandled tags recursively
            for child in element:
                ReportToolDocx.process_markdown_element(child, document)
            # Add any text directly within the unhandled tag as a paragraph
            text = ''.join(element.itertext()).strip() # Strip whitespace
            if text:
                 paragraph = document.add_paragraph()
                 paragraph.add_run(text)


    @staticmethod
    def _process_inline_elements(element, paragraph):
        """
        Helper function to process inline elements (bold, italics, text) within a block element
        and add them as runs to a python-docx paragraph or heading.
        """
        # Add text directly within the element before any children
        if element.text:
            paragraph.add_run(element.text)

        # Process children (inline elements)
        for child in element:
            if child.tag == 'strong':
                run = paragraph.add_run(''.join(child.itertext()))
                run.bold = True
            elif child.tag == 'em':
                run = paragraph.add_run(''.join(child.itertext()))
                run.italic = True
            elif child.tag == 'br':
                 # Add a line break within the paragraph
                 paragraph.add_run('\n')
            # Add handling for other inline elements like <code>, <a> etc. if needed
            else:
                 # For unhandled nested tags, just add their text content
                 paragraph.add_run(''.join(child.itertext()))

            # Add text that comes after a child element
            if child.tail:
                paragraph.add_run(child.tail)


    @staticmethod
    def make_report_docx(markdown_content: str, output_filename: str):
        """
        Converts Markdown content to DOCX preserving formatting.
        This is a static method as it operates solely on its inputs and doesn't require instance state.

        Args:
            markdown_content (str): The input string containing Markdown content.
            output_filename (str): The name of the output DOCX file.
        """
        # Check if output_filename is provided and is a string and ends with .docx
        if not isinstance(output_filename, str) or not output_filename.strip().lower().endswith('.docx'):
             print("Error: output_filename must be a non-empty string ending with .docx.")
             return

        # Create a new Document
        document = Document()

        # Add some default styles or modify existing ones if needed
        # Example: Set default font size for Normal style
        styles = document.styles
        style_normal = styles['Normal']
        font = style_normal.font
        font.name = 'Calibri' # Or your preferred font
        font.size = Pt(12)

        # Use markdown to convert the content to HTML
        # 'nl2br' extension converts newlines within paragraphs to <br> tags
        # 'extra' extension includes support for things like definition lists, footnotes, etc.
        # 'fenced_code' for code blocks
        # 'tables' for tables
        md = markdown.Markdown(extensions=['nl2br', 'extra', 'fenced_code', 'tables'])
        html_output = md.convert(markdown_content)

        # Wrap the HTML output in a root element to make it valid XML for parsing
        # Use a generic root tag like <root> instead of <div> to avoid potential conflicts
        wrapped_html = f"<root>{html_output}</root>"

        try:
            # Parse the HTML output into an element tree
            # ET.fromstring requires valid XML, so we wrap the HTML and use ET.XML
            root_element = ET.fromstring(wrapped_html)

            # Process the element tree and add elements to the document
            # The root element is the 'root' we added, process its children
            for element in root_element:
                # Call the static method using the class name
                ReportToolDocx.process_markdown_element(element, document)

        except ET.ParseError as e:
             print(f"Error parsing HTML output from Markdown: {e}")
             print("HTML Output:")
             print(wrapped_html)
             return # Stop if parsing fails
        except Exception as e:
            print(f"An unexpected error occurred during processing: {e}")
            return


        try:
            # Save the document
            document.save(output_filename)
            print(f"Successfully created {output_filename}")
        except Exception as e:
            print(f"Error saving DOCX file: {e}")


# Example Usage:
if __name__ == "__main__":
    example_markdown = """
# Report Title

This is the first paragraph. It contains some **bold text** and *italic text*. This line has a manual line break.
This is still the first paragraph.

## Section Heading

Here is a list:
* Item 1
* Item 2
    - Sub-item 1
    - Sub-item 2
+ Item 3

### Subsection

> This is a blockquote. (Note: Basic blockquote styling is not fully implemented, appears as a normal paragraph)

Another paragraph with a [link](https://example.com). (Note: Links are not clickable or specially formatted in this simple example)

---

End of the report.
"""

    # Call the static method using the class name, providing the filename explicitly
    ReportToolDocx.make_report_docx(example_markdown, "formatted_report_class.docx")
